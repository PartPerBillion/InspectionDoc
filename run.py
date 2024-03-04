import re
import pandas as pd
import docx
import numpy as np
import streamlit as st
from PIL import Image
import datetime
from docx.shared import Inches, Pt, Mm, RGBColor
from docx.oxml import OxmlElement, ns
from docxtpl import DocxTemplate

def line_space(doc, number_of_lines):
    for i in range(number_of_lines):
        p = doc.add_paragraph()
        p.style = None
        p.add_run(style = None)
        p.add_run(style = None)

def heading_number_generaator(heading, sub_heading, sub_sub_heading):
    if sub_heading == 0:
        return str(heading)
    elif sub_sub_heading == 0:
        return str(heading)+'.'+str(sub_heading)
    else:
        return str(heading)+'.'+str(sub_heading)+'.'+str(sub_sub_heading)
    
def add_heading_number(heading_number, level):
    headings = heading_number.split('.')
    if heading_number.count('.') == 2:
        heading = int(headings[0])
        sub_heading = int(headings[1])
        sub_sub_heading = int(headings[2])
    elif heading_number.count('.') == 1:
        heading = int(headings[0])
        sub_heading = int(headings[1])
        sub_sub_heading = 0
    else:
        heading = int(headings[0])
        sub_heading = 0
        sub_sub_heading = 0
    if level == 1:
        heading+=1
        return heading_number_generaator(heading,0,0)
    elif level == 2:
        sub_heading+=1
        return heading_number_generaator(heading,sub_heading,0)
    elif level == 3:
        sub_sub_heading+=1
        return heading_number_generaator(heading,sub_heading,sub_sub_heading)
    else:
        return(heading_number)
    
def removeEmptyColumns(df):
    for i in df.columns:
        if df[i].notnull().sum() < 5:
            df.drop(columns=i,inplace = True)
           
def add_points(x, doc):
    lines = x.split('\n')
    n = 0
    while n < len(lines):
        if len(lines[n]) > 0:
            if not lines[n].startswith(('-', '_')):
                for i in range(1, 6):
                    if n - i >= 0 and lines[n - i].startswith('-'):
                        lines[n - i] += lines[n]
                        del lines[n]
                        break
                    elif n - i >= 0 and lines[n - i].startswith('_'):
                        lines[n - i] += lines[n]
                        del lines[n]
                        break
            else:
                n += 1
        elif len(lines[n])<=1:
                del lines[n]
        else:
            n += 1
    for i in lines:
        if i.startswith('_'):
            add_paragraph_x(doc, i[2:] ,style = 'List Bullet', red = True)
        else:
            add_paragraph_x(doc, i[2:] ,style = 'List Bullet')
        # line_space(doc,1)

def add_site_observation_to_doc(string,doc,heading_number):
    rdict = {}
    n=0
    for i in string.split('\n'):
        if len(i)>0:
            if i[0] == '#':
                rdict[f'{n}_h'] = i[2:]
            elif i[0] == '$':
                rdict[f'{n}_s'] = i[2:]
            else:
                rdict[f'{n}_p'] = i
        n+=1
    for k,v in rdict.items():
        if len(v)>0:
            if k[-1] == 'h':
                heading_number = add_heading_number(heading_number,2)
                doc.add_heading(f'{heading_number}. {v.upper()}',2)
                # line_space(doc,1)
            elif k[-1] == 's':
                heading_number = add_heading_number(heading_number,3)
                doc.add_heading(f'{heading_number}. {v.upper()}',3)
                # line_space(doc,1)
            else:
                current_point = ''
                line = v.strip()
                if line.startswith('-'):
                    if current_point:
                        add_paragraph_x(doc, current_point.strip()[2:], style = 'List Bullet')
                        # line_space(doc,1)
                        current_point =  ''
                current_point += ' '+line
                
                if current_point:
                    add_paragraph_x(doc,current_point.strip()[2:], style = 'List Bullet')
                    # line_space(doc,1)
    return heading_number

def add_inspection_details_to_doc(text_file,doc,heading_number):
    rdict = {}
    n=0
    with open(text_file) as mytxt:
        for i in mytxt:
            if len(i)>0:
                if i[0] == '#':
                    rdict[f'{n}_h'] = i[2:]
                elif i[0] == '$':
                    rdict[f'{n}_s'] = i[2:]
                elif i[0] == '^':
                    rdict[f'{n}_t'] = i[2:].strip()
                elif i[0] == '%':
                    rdict[f'{n}_j'] = i[2:]
                elif i[0] == '>':
                    rdict[f'{n}_z'] = i[2:]
                else:
                    rdict[f'{n}_p'] = i
            n+=1
    for k,v in rdict.items():
        if len(v)>0:
            v = v.replace('\n','')
            if k[-1] == 'h':
                heading_number = add_heading_number(heading_number,1)
                doc.add_heading(f'{heading_number}. {v}'.upper(),1).bold = True
                # line_space(doc,1)
            elif k[-1] == 's':
                heading_number = add_heading_number(heading_number,2)
                doc.add_heading(f'{heading_number}. {v}'.upper(),2)
                # line_space(doc,1)
            elif k[-1] == 't':
                table = pd.read_csv(rf'Files/Inspection/TOWER INSPECTION BY ROBOTIC CRAWLER/{v}.csv')
                add_table_to_document(table, doc)
                line_space(doc,1)
            elif k[-1] == 'j':
                doc.add_picture(rf'Files/Inspection/TOWER INSPECTION BY ROBOTIC CRAWLER/{v}.jpg')
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                line_space(doc,1)
            elif k[-1] == 'z':
                p = doc.add_paragraph(v)
                p.style=doc.styles['Normal']
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                add_points(v,doc)
                    
    return heading_number

def table_of_contents(document):
    paragraph = document.add_paragraph()
    paragraph.style=document.styles['Normal']
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(ns.qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field"
    fldChar2.append(fldChar3)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(ns.qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    return document

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_table_to_document_fast(df, document, style = 'Table Grid',textStyle = 'Normal'):
    if len(df) > 0:
        table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = style
        for j, column in enumerate(df.columns):
            c = table.cell(0, j)
            p = c.paragraphs[0] 
            f = p.paragraph_format 
            f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(column, style = textStyle)
            r.bold = True
            c.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                c = table.cell(i + 1, j)
                p = c.paragraphs[0] 
                f = p.paragraph_format 
                f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                if type(value) == float:
                    r = p.add_run(str(round(value,2)), style = textStyle)
                else:
                    r = p.add_run(str(value), style = textStyle)
                c.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER      

def add_table_to_document(df, document, style = 'Table Grid'):
    if len(df) > 0:
        table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = style
        for j, column in enumerate(df.columns):
            cell = table.cell(0, j)
            
            p = table.cell(0, j).paragraphs[0]
            f = p.paragraph_format 
            f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(column)
            r.bold = True
            cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                cell = table.cell(i + 1, j)
                
                p = table.cell(i + 1, j).paragraphs[0]
                f = p.paragraph_format 
                f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(value))
                if str(value)[-1] == '%':
                    r.bold = True
                    if float(str(value)[:-2])>=20:
                        r.font.color.rgb = RGBColor(0xE3, 0x26, 0x36)
                    if float(str(value)[:-2])>=30:
                        r.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

def add_paragraph_x(doc,text,style,red = False):
    paragraph = doc.add_paragraph(style = style)
    # paragraph.style=doc.styles['Normal']
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    if red == True:
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        # Find all numbers with optional decimal places followed by %
        matches = re.finditer(r'\d+(\.\d{2,4})?%', text)

        # Keep track of the starting index in the original text
        current_index = 0

        for match in matches:
            start, end = match.span()
            number = match.group(0)
            # Add the preceding text before the matched number
            if start > current_index:
                run = paragraph.add_run(text[current_index:start])
                run.text = text[current_index:start]

            # Add the number in bold with a percentage sign
            bold_font = paragraph.add_run(number).font
            bold_font.bold = True
            current_index = end
            if float(number[:-1]) >= 20:
                bold_font.color.rgb = RGBColor(0xE3, 0x26, 0x36)
            if float(number[:-1]) >= 30:
                bold_font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

        # Add the remaining text after the last match
        if current_index < len(text):
            run = paragraph.add_run(text[current_index:])
            run.text = text[current_index:]

def create_paragraph_style(doc, name, size, font = 'Arial'):
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style(name, docx.enum.style.WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(size)
    obj_font.name = font

def add_table(doc, name, table, heading_number):
    if bool(table)==True:
            heading_number = add_heading_number(heading_number,1)
            doc.add_heading(f'{heading_number}. {name}'.upper(),1).bold = True
            if type(table) == list:
                for tables in table:
                    try:
                        overall_summary_df = pd.read_csv(tables,encoding='utf-8')
                    except:
                        overall_summary_df = pd.read_excel(tables)
                        pass
                    add_table_to_document(overall_summary_df,doc)
                    line_space(doc,1)
                doc.add_page_break()
            else:
                try:
                    df = pd.read_csv(table,encoding='utf-8')
                except:
                    df = pd.read_excel(table)
                add_table_to_document(df, doc)
                doc.add_page_break()

def add_first_page_header(doc, client_name, client_location, unit_number):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    footer = section.first_page_footer
    p =footer.add_paragraph()
    p.add_run('This document is the property of Aruntech Industrial Services Pvt Ltd, Chennai, India. It must not be stored, reproduced or disclosed to others without written authorisation from the Company. This document is subject to management review and update as deemed necessary.',style = 'SmallText')
    
    header = section.first_page_header 
    table = header.add_table(1, 3,width=Inches(7.2))
    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Inches(1.5)
    p = hdr_cells[0].paragraphs[0] 
    format = p.paragraph_format
    # format.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(r'Files/Client/Logo/ArunTech.jpg', width=Inches(.5))
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[0].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    hdr_cells[1].width = Inches(4)
    p = hdr_cells[1].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{client_name}\n', style = 'MediumText').bold = True
    p.add_run(f'{client_location} \n UNIT: {unit_number}').bold = True
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    hdr_cells[2].width = Inches(1.5)
    p = hdr_cells[2].paragraphs[0] 
    run = p.add_run()
    run.add_picture(rf'Files/Client/Logo/{client_name}.png', width=Inches(0.7))
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[2].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    table.style = 'Table Grid'
    line_space(header,1)

def add_header(doc, equipment_name, tag_number, unit_number, report_number, client_name, client_code, client_location, inspection_type):
    section = doc.sections[0]
    header = section.header 
    table = header.add_table(2, 3,width=Inches(7.2))
    hdr_cells = table.rows[0].cells
    hdr_cells[0].width = Inches(1.5)
    p = hdr_cells[0].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(r'Files/Client/Logo/ArunTech.jpg', width=Inches(.5))
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[0].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    hdr_cells[1].width = Inches(4.4)
    p = hdr_cells[1].paragraphs[0] 
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{equipment_name}\n{tag_number}\n UNIT: {unit_number}\n')
    f = r.font
    r.bold = True
    f.color.rgb = RGBColor(0xD0, 0x31, 0x2D)
    r = p.add_run(f'(Report No. {report_number})', style = 'SmallText')
    f = r.font
    r.bold = True
    f.color.rgb = RGBColor(0x00, 0x7B, 0xA7)
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
    hdr_cells[2].width = Inches(1.3)
    p = hdr_cells[2].paragraphs[0] 
    run = p.add_run()
    run.add_picture(rf'Files/Client/Logo/{client_name}.png', width=Inches(.7))
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[2].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

    sec_cells = table.rows[1].cells
    p = sec_cells[0].paragraphs[0] 
    p.add_run(f"Client:\n{client_code}-{unit_number}\n{client_location}")
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    sec_cells[0].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

    p = sec_cells[1].paragraphs[0] 
    p.add_run(f"{inspection_type}")
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    sec_cells[1].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

    p = sec_cells[2].paragraphs[0] 
    p.add_run('Page ')
    x = add_page_number(p.add_run())
    f = p.paragraph_format 
    f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    sec_cells[2].vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

    table.style = 'Table Grid'
    line_space(header,1)
def removeEmptyColumns(df):
    for i in df.columns:
        if df[i].notnull().sum() < 5:
            df.drop(columns=i,inplace = True)
# def add_detailed_report(file,doc):
#     excel = pd.ExcelFile(file)
#     for sheet in excel.sheet_names:
#         df = pd.read_excel(excel, sheet)
#         removeEmptyColumns(df)
#         cutoff = df.loc[df.iloc[:,0] == 'LOCATION'].index[0]
#         # try:
#         #     cutoff = df.loc[df['Unnamed: 0'] == 'LOCATION'].index[0]
#         # except:
#         #     cutoff = df.loc[df['Unnamed: 1'] == 'LOCATION'].index[0]
#         hf = df[:cutoff]
#         hf = hf[[hf.columns[0]]]
#         hf = hf.dropna(how='any',axis=0) 
#         hf = pd.DataFrame(hf[hf.columns[0]].str.split(':', n=1).tolist())
#         header = hf.iloc[0]
#         hf.columns = ['Info','Details']
#         hf = hf[1:]
#         hf['Info'] = hf['Info'].str.strip()
#         hf['Details'] = hf['Details'].str.strip()
#         bf = df[cutoff:]
#         header = bf.iloc[0].tolist()
#         headerStr = []
#         for i in header:
#             if str(i)[-8:] == '00:00:00':
#                 i = i.strftime('%b-%y')
#                 # i = str(i)[:-8]
#             headerStr.append(i)
        
#         bf.columns = headerStr
#         bf = bf[1:]
#         bf = bf.dropna(how='any',axis=0) 
#         for i in [hf,bf]:
#             i.reset_index(inplace=True)
#             i.drop(columns='index', inplace = True)

#         table1 = doc.add_table(rows=1, cols=1)
#         table1.style = 'Table Grid'
#         cell = table1.cell(0,0)
#         p = cell.paragraphs[0] 
#         f = p.paragraph_format 
#         f.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
#         r = p.add_run('Detailed Report')
#         r.bold = True
#         add_table_to_document_fast(hf,doc,textStyle='ReportText')
#         add_table_to_document_fast(bf,doc,textStyle='ReportText')
#         line_space(doc,1)
        
def make_inspection_document(client_name, client_location, unit_number, client_code, fpage_image, inspection_date, equipment_name, tag_number, inspection_type, edited_df, result_and_conclusion, site_observation, overall_summary, thickness_details, scanning_details, shellwise_inspection, tower_drawing, shell_plate_pics):
    doc = docx.Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = docx.shared.Pt(11)
    create_paragraph_style(doc, 'CommentsStyle', 16)
    create_paragraph_style(doc, 'BigText', 36)
    create_paragraph_style(doc, 'MediumText', 14)
    create_paragraph_style(doc, 'ReportText', 10)
    create_paragraph_style(doc, 'SmallText', 8)

    add_first_page_header(doc, client_name, client_location, unit_number)
    

    table_0 = doc.add_table(1,2)
    p = table_0.rows[0].cells[0].paragraphs[0]
    format = p.paragraph_format
    inspection_date = pd.to_datetime(inspection_date)
    date_str = inspection_date.strftime('%B %Y/%d')
    date_components = date_str.split(' ')
    date_components[0] = date_components[0][:3]
    date_components[0] = date_components[0].upper()
    datex = ' '.join(date_components)
    report_number = f'ATL/UT/{client_code}-{client_location}/{datex}'
    p.add_run('Report no:\n', style = 'SmallText')
    p.add_run(f"{report_number}")
    p = table_0.rows[0].cells[1].paragraphs[0]
    format = p.paragraph_format

    inspection_dateX = inspection_date.strftime('%d-%m-%Y')
    p.add_run('Inspection date:\n',style='SmallText')
    p.add_run(f'{inspection_dateX}')
    
    table_0.style = 'Medium Shading 1 Accent 6'
  
    table_1 = doc.add_table(2,1)
    p = table_1.rows[0].cells[0].paragraphs[0]
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{equipment_name}', style = 'CommentsStyle').bold = False
    p.add_run(f'\n{tag_number}',style = 'CommentsStyle').bold = True
    p = table_1.rows[1].cells[0].paragraphs[0]
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{inspection_type}', style = 'CommentsStyle').bold = False
    
    table_1.style = 'Medium Shading 1 Accent 6'

    line_space(doc,4)
    if bool(fpage_image)==True:
        image = Image.open(fpage_image)
        image.save(rf'Files/Temp/img.png')
        doc.add_picture(rf'Files/Temp/img.png', height = Inches(3.5))
    else:
        doc.add_picture(rf'Files/Inspection/{inspection_type}/FrontPageImage.jpg', height = Inches(3.5))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    line_space(doc,4)
    edited_df.to_csv(rf'Files/Temp/authors.csv')
    authors = pd.read_csv(rf'Files/Temp/authors.csv')
    authors['Date'] = pd.to_datetime(authors['Date'])
    authors['Date'] = authors['Date'].dt.strftime('%d-%m-%Y')
    authors.drop('Unnamed: 0', inplace=True, axis=1)
    add_table_to_document(authors, doc, 'Medium Shading 1 Accent 6')


    add_header(doc, equipment_name, tag_number, unit_number, report_number, client_name, client_code, client_location, inspection_type)

    p = doc.add_paragraph()
    p.add_run('Table of contents'.upper(), style = 'MediumText').bold = True
    line_space(doc,1)
    format = p.paragraph_format
    format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    doc = table_of_contents(doc)
    doc.add_page_break()
    heading_number = '1'
    doc.add_heading(f'{heading_number}. Results and conclusion'.upper(),1).bold = True

    paragraph_format = doc.styles['List Bullet'].paragraph_format
    paragraph_format.left_indent = Inches(0.5)

    add_points(result_and_conclusion, doc)
    doc.add_page_break()

    heading_number = add_heading_number(heading_number,1)
    doc.add_heading(f'{heading_number}. Site observation'.upper(),1).bold = True
    line_space(doc,1)
    add_site_observation_to_doc(site_observation,doc,heading_number)
    doc.add_page_break()

    add_table(doc, 'Overall Summary', overall_summary, heading_number)
    add_table(doc, 'Thickness Details', thickness_details, heading_number)
    add_table(doc, 'Scanning Details', scanning_details, heading_number)
    add_table(doc, 'Shellwise Inspection', shellwise_inspection, heading_number)


    if bool(tower_drawing)==True:
        heading_number = add_heading_number(heading_number,1)
        doc.add_heading(f'{heading_number}. Tower Drawings'.upper(),1).bold = True
        for pic in tower_drawing:
            image = Image.open(pic)
            image.save(rf'Files/Temp/img.png')
            line_space(doc,2)
            doc.add_picture(rf'Files/Temp/img.png', height = Inches(5))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            line_space(doc,1)
        doc.add_page_break()
        

    if bool(shell_plate_pics)==True:
        heading_number = add_heading_number(heading_number,1)
        doc.add_heading(f'{heading_number}. Shellplate Pictures'.upper(),1).bold = True
        p = doc.add_paragraph()
        r = p.add_run()
        for pic in shell_plate_pics:
            image = Image.open(pic)
            image.save(rf'Files/Temp/img.png')
            r.add_picture(rf'Files/Temp/img.png', width = Inches(2.2))
            r.add_text(' ')
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        line_space(doc,1)

    heading_number = add_heading_number(heading_number,1)
    doc.add_heading(f'{heading_number}. Detailed Report'.upper(),1).bold = True
    doc.add_page_break()
    
    # if len(list(detailed_report.keys())[0])>0:
    #     if bool(list(detailed_report.values())[0])==True:
    #         heading_number = add_heading_number(heading_number,1)
    #         doc.add_heading(f'{heading_number}. Detailed Report'.upper(),1).bold = True
    #         for k,v in detailed_report.items():
    #             heading_number = add_heading_number(heading_number,2)
    #             doc.add_heading(f'{heading_number}. {k.upper()}',2)
    #             line_space(doc,4)
    #             p = doc.add_paragraph()
    #             p.add_run(f'{v[0].upper()}', style = 'BigText').bold = True
    #             format = p.paragraph_format
    #             format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    #             doc.add_page_break()
    #             # for table in v[1]:
    #             # try:
    #             #     overall_summary_df = pd.read_csv(table)
    #             #     # add_detailed_report(table,doc)
    #             # except:
    #             # s(v[1])
    #             add_detailed_report(v[1],doc)
    #             # overall_summary_df = pd.read_excel(overall_summary)
                

    #                 # add_table_to_document_fast(overall_summary_df,doc)
    #                 # line_space(doc,1)
    #             doc.add_page_break()

    add_inspection_details_to_doc(rf'Files/Inspection/{inspection_type}/text.txt',doc,heading_number)

    return doc
